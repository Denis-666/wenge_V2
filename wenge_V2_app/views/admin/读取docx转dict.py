
from docx import Document

import os


def docxToDict(FileNameFullPath) -> dict :
    print('目前运行:  docxToDict(FileNameFullPath) -> dict :')
    document = Document(FileNameFullPath)
    paragraphs = document.paragraphs

    dict_0 = {'sucess':{},'fail':{}}
    id_sucess = 1
    id_fail = 1
    for paragraph in document.paragraphs:
        print(paragraph.text.split('----'))
        each_line = paragraph.text
        try:
            str_3 = each_line.split('----')
            账号 = str_3[0]
            密码 = str_3[1]
            try:
                extraEmail= str_3[2]
            except:
                extraEmail = None
            print('id:', id_sucess, '账号:', 账号, '密码:', 密码,'extraEmail',extraEmail)
            dict_0['sucess'][id_sucess] = {'username': 账号, 'password': 密码,'extraEmail':extraEmail}
            id_sucess = id_sucess + 1
        except:
            # 输出错误信息 (没写好)
            
            dict_0['fail'][id_fail] = {'line':each_line}
            print(id_fail, '解析失败')
            id_fail = id_fail + 1

    return dict_0
    


